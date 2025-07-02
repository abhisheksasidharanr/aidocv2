from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
import requests
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from docx.shared import RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter


app = Flask(__name__)
CORS(app)

API_KEY = "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"
API_URL = "https://api.groq.com/openai/v1/chat/completions"

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    extracted_text = []

    # Extract text from paragraphs (normal document text)
    for para in doc.paragraphs:
        extracted_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            extracted_text.append("\t".join(row_text))  # Tab-separated values for readability

    return "\n".join(extracted_text)  # Join with new lines to maintain structure


def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])


def compare_documents(original_text, edited_text):
    prompt = f"""Compare the following **original text** and **edited text**, marking the differences at both **word and letter levels**.

    **Rules:**
    1. **Highlight deleted words as** `<DEL>word</DEL>`
    2. **Highlight added words as** `<ADD>word</ADD>`
    3. **Highlight deleted letters inside words as** `<DEL>letter</DEL>`
    4. **Highlight added letters inside words as** `<ADD>letter</ADD>`
    5. **Preserve line breaks, punctuation, and spaces exactly as in the edited text.**
    6. **Return only the compared text with markers—no explanations or extra formatting.**
    7. **If `</DEL><ADD>` appears in the output, insert a space between them (`</DEL> <ADD>`).**
    8. **If the original and edited texts contain a table, preserve its structure and format it using <table>, <tr>, and <td> HTML tags, ensuring that added or deleted content is appropriately highlighted within table cells.**

    ---

    **Example 1**  
        **Original:** `Hi I am Verbat.`  
        **Edited:** `Hi I am Abhishek. I am from Verbat.`  
        **Output:**  
    `Hi I am <ADD>Abhishek</ADD> <DEL>Verbat</DEL>. <ADD>I am from Verbat.</ADD>`

    ---

    **Example 2**  
        **Original:** `color`  
        **Edited:** `colour`  
        **Output:**  
    `col<ADD>o</ADD>our`

    ---

    Now, compare the following:

    **Original:**  
    {original_text}

    **Edited:**  
    {edited_text}

    **Output:**
"""

    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    data = {"model": "qwen/qwen3-32b", "messages": [{"role": "system", "content": prompt}], "temperature": 0.2}

    response = requests.post(API_URL, headers=headers, json=data)
    return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    # return response


@app.route("/compare", methods=["POST"])
def compare_files():
    if "original" not in request.files or "edited" not in request.files:
        return jsonify({"error": "Both original and edited files are required"}), 400

    original = request.files["original"]
    edited = request.files["edited"]

    original_path = os.path.join(UPLOAD_FOLDER, original.filename)
    edited_path = os.path.join(UPLOAD_FOLDER, edited.filename)

    original.save(original_path)
    edited.save(edited_path)

    if original.filename.endswith(".docx"):
        original_text = extract_text_from_docx(original_path)
        edited_text = extract_text_from_docx(edited_path)
    elif original.filename.endswith(".pdf"):
        original_text = extract_text_from_pdf(original_path)
        edited_text = extract_text_from_pdf(edited_path)
    else:
        return jsonify({"error": "Unsupported file format"}), 400

    diff_result = compare_documents(original_text, edited_text)

    
    # generate_highlighted_pdf(diff_result, output_pdf_path)

    return jsonify({
        "original_text": original_text,
        "edited_text": edited_text,
        "result" : diff_result,
        "hii" : "stop"
    })





if __name__ == "__main__":
    app.run(debug=True, port=5000)
